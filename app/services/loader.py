# === app/services/loader.py ===
import os
import tempfile
import requests
import mimetypes
import docx
import csv
import pandas as pd
from bs4 import BeautifulSoup
from email import policy
from email import message_from_binary_file
import pdfplumber
from PIL import Image
import pytesseract
import fitz  # PyMuPDF for PDF image extraction
import filetype  # replaces python-magic
import zipfile
from io import BytesIO

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".txt", ".csv", ".xlsx", ".eml",
    ".jpg", ".jpeg", ".png", ".tiff", ".bin"
}

# ---------- File Download ----------
def download_file(url: str) -> str:
    url = str(url)
    if "drive.google.com" in url:
        file_id = None
        if "/d/" in url:
            file_id = url.split("/d/")[1].split("/")[0]
        elif "id=" in url:
            file_id = url.split("id=")[1].split("&")[0]
        if file_id:
            url = f"https://drive.google.com/uc?export=download&id={file_id}"
    response = requests.get(url, allow_redirects=True)
    response.raise_for_status()
    ext = guess_extension(url, response)
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(response.content)
        return tmp.name

def guess_extension(url: str, response) -> str:
    if "." in url.split("?")[0]:
        ext = "." + url.split("?")[0].split(".")[-1].lower()
        if ext in SUPPORTED_EXTENSIONS:
            return ext
    ctype = response.headers.get("Content-Type")
    if ctype:
        ext = mimetypes.guess_extension(ctype.split(";")[0].strip())
        if ext:
            return ext
    return ".tmp"

# ---------- Extractors ----------

def extract_text_from_pdf(file_path: str) -> str:
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
            else:
                pil_img = page.to_image(resolution=300).original
                ocr_text = pytesseract.image_to_string(pil_img)
                if ocr_text.strip():
                    text.append(ocr_text)
    return "\n".join(text)

def extract_images_from_pdf(file_path: str) -> list[Image.Image]:
    images = []
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)
        for img in image_list:
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            pil_img = Image.open(BytesIO(image_bytes)).convert("RGB")
            images.append(pil_img)
    return images

def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_images_from_docx(file_path: str) -> list[Image.Image]:
    document = docx.Document(file_path)
    images = []
    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            pil_img = Image.open(BytesIO(image_bytes)).convert("RGB")
            images.append(pil_img)
    return images

def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_csv(file_path: str) -> str:
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as csvfile:
        return "\n".join([" | ".join(row) for row in csv.reader(csvfile)])

def extract_text_from_xlsx(file_path: str) -> str:
    dfs = pd.read_excel(file_path, sheet_name=None)
    return "\n".join(df.to_csv(index=False) for df in dfs.values())

def extract_text_from_image(file_path: str) -> str:
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)

def extract_text_from_email(file_path: str) -> str:
    texts = []
    with open(file_path, "rb") as f:
        msg = message_from_binary_file(f, policy=policy.default)
    for part in msg.walk():
        content_disposition = part.get_content_disposition()
        content_type = part.get_content_type()
        if content_disposition == "attachment":
            filename = part.get_filename()
            if filename:
                suffix = os.path.splitext(filename)[-1].lower()
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(part.get_payload(decode=True))
                    texts.append(load_from_path(tmp.name)[0])
        elif content_type in ["text/plain", "text/html"]:
            payload = part.get_payload(decode=True)
            if payload:
                if content_type == "text/html":
                    soup = BeautifulSoup(payload, "html.parser")
                    texts.append(soup.get_text())
                else:
                    texts.append(payload.decode(errors="ignore"))
    return "\n".join(filter(None, texts))

# ---------- Advanced .bin Handling ----------

def detect_file_type(file_path: str) -> str:
    kind = filetype.guess(file_path)
    if kind:
        return kind.mime
    return ""

def extract_files_from_zip(file_path: str) -> list[str]:
    extracted_paths = []
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(tmpdir)
                for root, _, files in os.walk(tmpdir):
                    for file in files:
                        extracted_paths.append(os.path.join(root, file))
        except zipfile.BadZipFile:
            pass
        return extracted_paths

def extract_text_from_bin(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception:
            return ""
    except Exception:
        return ""

def extract_text_and_images_from_bin(file_path: str):
    text = ""
    images = []
    file_type = detect_file_type(file_path)

    if file_type == "application/pdf":
        text = extract_text_from_pdf(file_path)
        images = extract_images_from_pdf(file_path)
    elif file_type and file_type.startswith("image/"):
        images = [Image.open(file_path).convert("RGB")]
    elif file_type == "application/zip":
        extracted_files = extract_files_from_zip(file_path)
        for ef in extracted_files:
            t, i = load_from_path(ef)
            text += t + "\n"
            images.extend(i)
    else:
        text = extract_text_from_bin(file_path)

    return text, images

# ---------- Main Loader ----------

def load_from_path(path: str):
    ext = os.path.splitext(path)[-1].lower()
    text = ""
    images = []

    if ext == ".pdf":
        text = extract_text_from_pdf(path)
        images = extract_images_from_pdf(path)
    elif ext == ".docx":
        text = extract_text_from_docx(path)
        images = extract_images_from_docx(path)
    elif ext == ".txt":
        text = extract_text_from_txt(path)
    elif ext == ".csv":
        text = extract_text_from_csv(path)
    elif ext == ".xlsx":
        text = extract_text_from_xlsx(path)
    elif ext in {".jpg", ".jpeg", ".png", ".tiff"}:
        images = [Image.open(path).convert("RGB")]
    elif ext == ".eml":
        text = extract_text_from_email(path)
    elif ext == ".bin":
        text, images = extract_text_and_images_from_bin(path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    return text, images


def load_from_url(url: str):
    path = download_file(url)
    return load_from_path(path)
